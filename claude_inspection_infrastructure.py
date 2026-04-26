#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════╗
║         Ceph 集群自动化巡检脚本 v1.0                         ║
║  适配: Pacific 16.2.x | cephadm + Docker 部署方式            ║
║  支持: 多集群并行巡检 | Word 格式巡检报告                     ║
╚══════════════════════════════════════════════════════════════╝
 
依赖安装:
    pip install requests python-docx
 
使用方法:
    python3 ceph_inspector.py
    python3 ceph_inspector.py --output /tmp/reports
    python3 ceph_inspector.py --cluster 集群-01   # 只巡检指定集群
"""
 
import argparse
import datetime
import os
import sys
import warnings
from typing import Dict, List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
 
import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
warnings.filterwarnings("ignore", message="Unverified HTTPS request")
 
# ═══════════════════════════════════════════════════════════════
#  ⚙  配置区域 — 请根据实际环境修改以下内容
# ═══════════════════════════════════════════════════════════════
 
CLUSTERS: List[Dict] = [
    {
        "name":           "Ceph集群-01",
        "env":            "生产环境",
        "description":    "fs-1f8hh0104 机房 / 主存储集群",
        "version":        "16.2.11",
        # cephadm 部署 Prometheus 默认端口 9095；若自定义请修改
        "prometheus_url": "http://192.168.x.1:9095",
        # 若 Prometheus 需要 Basic Auth，填写用户名/密码，否则留空
        "prom_user":      "",
        "prom_pass":      "",
    },
    {
        "name":           "Ceph集群-02",
        "env":            "生产环境",
        "description":    "集群-02 / 备份存储",
        "version":        "16.2.13",
        "prometheus_url": "http://192.168.x.2:9095",
        "prom_user":      "",
        "prom_pass":      "",
    },
    {
        "name":           "Ceph集群-03",
        "env":            "生产环境",
        "description":    "集群-03 / 对象存储",
        "version":        "16.2.15",
        "prometheus_url": "http://192.168.x.3:9095",
        "prom_user":      "",
        "prom_pass":      "",
    },
    {
        "name":           "Ceph集群-04",
        "env":            "生产环境",
        "description":    "集群-04 / 块存储",
        "version":        "16.2.15",
        "prometheus_url": "http://192.168.x.4:9095",
        "prom_user":      "",
        "prom_pass":      "",
    },
]
 
REQUEST_TIMEOUT = 15          # 秒；网络慢可适当调大
REPORT_OUTPUT_DIR = "./reports"
 
# ═══════════════════════════════════════════════════════════════
#  PromQL 查询封装
# ═══════════════════════════════════════════════════════════════
 
def _prom_get(url: str, endpoint: str, params: dict,
              user: str, passwd: str) -> Optional[dict]:
    auth = (user, passwd) if user else None
    try:
        r = requests.get(
            f"{url}{endpoint}", params=params,
            timeout=REQUEST_TIMEOUT, verify=False, auth=auth,
        )
        r.raise_for_status()
        return r.json()
    except requests.exceptions.ConnectionError:
        print(f"    [ERROR] 无法连接 {url} ，请确认 Prometheus 地址和端口")
    except requests.exceptions.Timeout:
        print(f"    [ERROR] 连接超时: {url}")
    except Exception as exc:
        print(f"    [ERROR] 请求异常: {exc}")
    return None
 
 
def query_scalar(url: str, query: str, user="", passwd="") -> Optional[float]:
    """返回单个数值，查不到返回 None"""
    data = _prom_get(url, "/api/v1/query", {"query": query}, user, passwd)
    if data and data.get("status") == "success":
        result = data["data"]["result"]
        if result:
            try:
                return float(result[0]["value"][1])
            except (KeyError, IndexError, ValueError):
                pass
    return None
 
 
def query_vector(url: str, query: str, user="", passwd="") -> List[Dict]:
    """返回 instant vector 结果列表"""
    data = _prom_get(url, "/api/v1/query", {"query": query}, user, passwd)
    if data and data.get("status") == "success":
        return data["data"]["result"]
    return []
 
 
def query_alerts(url: str, user="", passwd="") -> List[Dict]:
    """从 Alertmanager API 获取活跃告警"""
    # 先尝试 Prometheus /api/v1/alerts
    data = _prom_get(url, "/api/v1/alerts", {}, user, passwd)
    if data and data.get("status") == "success":
        alerts = []
        for a in data["data"].get("alerts", []):
            if a.get("state") == "firing":
                alerts.append({
                    "name":     a["labels"].get("alertname", "-"),
                    "severity": a["labels"].get("severity", "-"),
                    "summary":  a.get("annotations", {}).get("summary", "-"),
                    "state":    a.get("state", "-"),
                    "fired_at": a.get("activeAt", "")[:19].replace("T", " "),
                })
        return alerts
    return []
 
# ═══════════════════════════════════════════════════════════════
#  数据采集器
# ═══════════════════════════════════════════════════════════════
 
class CephCollector:
    def __init__(self, cfg: Dict):
        self.cfg  = cfg
        self.url  = cfg["prometheus_url"].rstrip("/")
        self.user = cfg.get("prom_user", "")
        self.pwd  = cfg.get("prom_pass", "")
 
    def _q(self, query: str) -> Optional[float]:
        return query_scalar(self.url, query, self.user, self.pwd)
 
    def _v(self, query: str) -> List[Dict]:
        return query_vector(self.url, query, self.user, self.pwd)
 
    def _metric_key(self, item: Dict) -> str:
        """从 metric labels 中提取 daemon/osd/mon 等标识"""
        m = item.get("metric", {})
        for k in ("ceph_daemon", "osd", "mon", "pool_id", "name", "instance"):
            if k in m:
                return m[k]
        return list(m.values())[-1] if m else ""
 
    def collect(self) -> Dict:
        print(f"\n  ► 采集 [{self.cfg['name']}]  {self.url}")
        return {
            "cluster_info": self.cfg,
            "collected_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "reachable":    True,
            "health":       self._health(),
            "capacity":     self._capacity(),
            "osd":          self._osd(),
            "mon":          self._mon(),
            "pg":           self._pg(),
            "pool":         self._pool(),
            "performance":  self._performance(),
            "alerts":       self._alerts(),
        }
 
    # ── 健康状态 ──────────────────────────────────────────────
    def _health(self) -> Dict:
        val = self._q("ceph_health_status")
        mp = {0: ("HEALTH_OK", "正常"), 1: ("HEALTH_WARN", "告警"), 2: ("HEALTH_ERR", "错误")}
        if val is None:
            return {"value": None, "code": "UNKNOWN", "cn": "无法获取"}
        code, cn = mp.get(int(val), ("UNKNOWN", "未知"))
        return {"value": val, "code": code, "cn": cn}
 
    # ── 容量 ──────────────────────────────────────────────────
    def _capacity(self) -> Dict:
        total = self._q("ceph_cluster_total_bytes")
        used  = self._q("ceph_cluster_total_used_bytes")
        avail = (total - used) if (total and used) else None
        pct   = (used / total * 100) if (total and used and total > 0) else None
        return {"total": total, "used": used, "avail": avail, "usage_pct": pct}
 
    # ── OSD ───────────────────────────────────────────────────
    def _osd(self) -> Dict:
        up_vec    = self._v("ceph_osd_up")
        in_vec    = self._v("ceph_osd_in")
        bytes_vec = self._v("ceph_osd_stat_bytes")
        used_vec  = self._v("ceph_osd_stat_bytes_used")
        apply_vec = self._v("ceph_osd_apply_latency_ms")
        commit_vec= self._v("ceph_osd_commit_latency_ms")
 
        def vec_map(vec):
            return {self._metric_key(i): float(i["value"][1]) for i in vec}
 
        up_m  = vec_map(up_vec)
        in_m  = vec_map(in_vec)
        b_m   = vec_map(bytes_vec)
        u_m   = vec_map(used_vec)
 
        all_ids = sorted(set(up_m) | set(in_m), key=lambda x: (len(str(x)), str(x)))
        details = []
        for oid in all_ids:
            tb = b_m.get(oid, 0)
            ub = u_m.get(oid, 0)
            details.append({
                "id":          oid,
                "up":          int(up_m.get(oid, 0)),
                "in":          int(in_m.get(oid, 0)),
                "total_bytes": tb,
                "used_bytes":  ub,
                "usage_pct":   (ub / tb * 100) if tb > 0 else 0,
            })
 
        total = len(details)
        up_cnt = sum(1 for d in details if d["up"] == 1)
        in_cnt = sum(1 for d in details if d["in"] == 1)
 
        def avg_lat(vec):
            vals = [float(i["value"][1]) for i in vec if float(i["value"][1]) > 0]
            return sum(vals) / len(vals) if vals else None
 
        return {
            "total":   total,
            "up":      up_cnt,
            "in":      in_cnt,
            "down":    total - up_cnt,
            "out":     total - in_cnt,
            "details": details,
            "avg_apply_lat_ms":  avg_lat(apply_vec),
            "avg_commit_lat_ms": avg_lat(commit_vec),
        }
 
    # ── MON ───────────────────────────────────────────────────
    def _mon(self) -> Dict:
        vec = self._v("ceph_mon_quorum_status")
        details = [
            {"name": self._metric_key(i), "in_quorum": int(float(i["value"][1]))}
            for i in vec
        ]
        total    = len(details)
        in_q     = sum(1 for d in details if d["in_quorum"] == 1)
        return {"total": total, "in_quorum": in_q, "out_quorum": total - in_q, "details": details}
 
    # ── PG ────────────────────────────────────────────────────
    def _pg(self) -> Dict:
        keys = ["total", "active", "clean", "degraded", "recovering",
                "recovery_wait", "undersized", "stale", "peering",
                "remapped", "backfill_wait", "forced_recovery"]
        result = {}
        for k in keys:
            result[k] = self._q(f"ceph_pg_{k}")
        return result
 
    # ── Pool ──────────────────────────────────────────────────
    def _pool(self) -> Dict:
        def pool_map(vec):
            m = {}
            for i in vec:
                met = i.get("metric", {})
                pid = met.get("pool_id", met.get("name", ""))
                name = met.get("name", "")
                m[pid] = {"name": name, "val": float(i["value"][1])}
            return m
 
        used_m    = pool_map(self._v("ceph_pool_bytes_used"))
        avail_m   = pool_map(self._v("ceph_pool_max_avail"))
        obj_m     = pool_map(self._v("ceph_pool_objects"))
        rd_m      = pool_map(self._v("irate(ceph_pool_rd[5m])"))
        wr_m      = pool_map(self._v("irate(ceph_pool_wr[5m])"))
        rd_bm     = pool_map(self._v("irate(ceph_pool_rd_bytes[5m])"))
        wr_bm     = pool_map(self._v("irate(ceph_pool_wr_bytes[5m])"))
 
        all_pids = sorted(
            set(used_m) | set(avail_m) | set(obj_m),
            key=lambda x: (len(str(x)), str(x))
        )
        pools = []
        for pid in all_pids:
            used  = used_m.get(pid,  {}).get("val", 0)
            avail = avail_m.get(pid, {}).get("val", 0)
            pname = (used_m.get(pid) or avail_m.get(pid) or {}).get("name", f"pool-{pid}")
            total = used + avail
            pools.append({
                "id":         pid,
                "name":       pname,
                "used":       used,
                "avail":      avail,
                "total":      total,
                "usage_pct":  (used / total * 100) if total > 0 else 0,
                "objects":    obj_m.get(pid,  {}).get("val", 0),
                "rd_iops":    rd_m.get(pid,   {}).get("val", 0),
                "wr_iops":    wr_m.get(pid,   {}).get("val", 0),
                "rd_bytes_s": rd_bm.get(pid,  {}).get("val", 0),
                "wr_bytes_s": wr_bm.get(pid,  {}).get("val", 0),
            })
        return {"pools": pools}
 
    # ── 性能 ──────────────────────────────────────────────────
    def _performance(self) -> Dict:
        rd_iops = self._q("sum(irate(ceph_osd_op_r[5m]))")
        wr_iops = self._q("sum(irate(ceph_osd_op_w[5m]))")
        rd_bps  = self._q("sum(irate(ceph_osd_op_r_out_bytes[5m]))")
        wr_bps  = self._q("sum(irate(ceph_osd_op_w_in_bytes[5m]))")
 
        # 延迟：sum/count 相除
        def lat_ms(sum_q, cnt_q):
            s = self._q(sum_q)
            c = self._q(cnt_q)
            if s is not None and c and c > 0:
                return s / c * 1000
            return None
 
        rd_lat = lat_ms("sum(irate(ceph_osd_op_r_latency_sum[5m]))",
                         "sum(irate(ceph_osd_op_r_latency_count[5m]))")
        wr_lat = lat_ms("sum(irate(ceph_osd_op_w_latency_sum[5m]))",
                         "sum(irate(ceph_osd_op_w_latency_count[5m]))")
 
        return {
            "rd_iops":    rd_iops,
            "wr_iops":    wr_iops,
            "rd_bytes_s": rd_bps,
            "wr_bytes_s": wr_bps,
            "rd_lat_ms":  rd_lat,
            "wr_lat_ms":  wr_lat,
        }
 
    # ── 告警 ──────────────────────────────────────────────────
    def _alerts(self) -> List[Dict]:
        return query_alerts(self.url, self.user, self.pwd)
 
 
# ═══════════════════════════════════════════════════════════════
#  格式化工具
# ═══════════════════════════════════════════════════════════════
 
def bh(b: Optional[float], d: int = 2) -> str:
    """字节转人类可读"""
    if b is None:
        return "N/A"
    b = float(b)
    for u in ("B", "KB", "MB", "GB", "TB", "PB"):
        if abs(b) < 1024.0:
            return f"{b:.{d}f} {u}"
        b /= 1024.0
    return f"{b:.{d}f} EB"
 
 
def fp(v: Optional[float], d: int = 2) -> str:
    return f"{v:.{d}f}%" if v is not None else "N/A"
 
 
def fv(v: Optional[float], suf: str = "", d: int = 2) -> str:
    return f"{v:.{d}f}{suf}" if v is not None else "N/A"
 
 
HEALTH_COLOR = {"HEALTH_OK": "00B050", "HEALTH_WARN": "FFC000",
                "HEALTH_ERR": "FF0000", "UNKNOWN": "808080"}
 
THEME_BLUE  = "1F497D"
THEME_DARK  = "2E4057"
THEME_RED   = "C00000"
THEME_GRAY  = "F2F2F2"
THEME_GREEN = "E2EFDA"
THEME_YELL  = "FFF2CC"
 
# ═══════════════════════════════════════════════════════════════
#  Word 报告工具函数
# ═══════════════════════════════════════════════════════════════
 
def _cell_bg(cell, color: str):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pr.append(shd)
 
 
def _cell_border(cell):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    pr.append(borders)
 
 
def _set_font(run, size=9, bold=False, color: Optional[Tuple] = None):
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass
 
 
def _para_font(para):
    for run in para.runs:
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        except Exception:
            pass
 
 
def tbl_header(table, headers: List[str], bg: str = THEME_BLUE):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        _set_font(run, size=9, bold=True, color=(0xFF, 0xFF, 0xFF))
        _cell_bg(cell, bg)
        _cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
 
 
def tbl_row(table, row_idx: int, values: List[str],
            center_cols: List[int] = None, alt: bool = False,
            row_color: str = None):
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        if center_cols and i in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val) if val is not None else "-")
        _set_font(run, size=9)
        bg = row_color or (THEME_GRAY if alt else "FFFFFF")
        _cell_bg(cell, bg)
        _cell_border(cell)
 
 
def add_h(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    _para_font(p)
    return p
 
 
def add_p(doc: Document, text: str, bold=False, size=10, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p
 
 
def make_table(doc: Document, headers: List[str], header_bg: str = THEME_BLUE) -> "Table":
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_header(table, headers, header_bg)
    return table
 
 
def add_row(table, values: List[str], center_cols=None, alt=False, row_color=None):
    row = table.add_row()
    row_idx = len(table.rows) - 1
    tbl_row(table, row_idx, values, center_cols, alt, row_color)
 
 
# ═══════════════════════════════════════════════════════════════
#  巡检建议生成
# ═══════════════════════════════════════════════════════════════
 
def make_recommendations(data: Dict) -> List[Tuple[str, str]]:
    """返回 (类别标签, 建议内容) 元组列表"""
    recs = []
    h    = data["health"]
    cap  = data["capacity"]
    osd  = data["osd"]
    mon  = data["mon"]
    pg   = data["pg"]
    perf = data["performance"]
    alts = data["alerts"]
 
    # 健康
    if h["code"] == "HEALTH_OK":
        recs.append(("✅ 健康状态", "集群整体健康（HEALTH_OK），无需处理。"))
    elif h["code"] == "HEALTH_WARN":
        recs.append(("⚠️ 健康状态", "集群存在告警（HEALTH_WARN），建议执行 `ceph health detail` 获取详细信息并逐项排查。"))
    elif h["code"] == "HEALTH_ERR":
        recs.append(("❌ 健康状态", "集群处于错误状态（HEALTH_ERR），需要立即处理！请检查 OSD/MON 及 PG 状态。"))
    else:
        recs.append(("⚠️ 健康状态", "无法获取集群健康状态，请检查 Prometheus 连通性。"))
 
    # 容量
    pct = cap.get("usage_pct")
    if pct is None:
        recs.append(("⚠️ 存储容量", "容量数据获取失败，请检查 Prometheus 指标 ceph_cluster_total_bytes。"))
    elif pct >= 85:
        recs.append(("❌ 存储容量", f"容量使用率已达 {pct:.1f}%，超过 85% 安全水位，需立即制定扩容方案！"))
    elif pct >= 70:
        recs.append(("⚠️ 存储容量", f"容量使用率 {pct:.1f}%，已超 70% 预警线，请关注增长趋势并提前规划扩容。"))
    else:
        recs.append(("✅ 存储容量", f"容量使用率 {pct:.1f}%，处于健康范围，继续关注趋势。"))
 
    # OSD
    if osd["down"] > 0:
        recs.append(("❌ OSD 状态", f"检测到 {osd['down']} 个 OSD 处于 DOWN 状态，请立即检查对应节点磁盘健康状态及系统日志（journalctl -u ceph-osd@<id>）。"))
    if osd["out"] > 0:
        recs.append(("⚠️ OSD 状态", f"有 {osd['out']} 个 OSD 处于 OUT 状态，集群正进行数据回填，IO 性能可能受影响，请持续关注 PG 恢复进度。"))
    if osd["down"] == 0 and osd["out"] == 0:
        recs.append(("✅ OSD 状态", f"所有 {osd['total']} 个 OSD 均处于 UP/IN 状态，运行正常。"))
 
    apply_lat  = osd.get("avg_apply_lat_ms")
    commit_lat = osd.get("avg_commit_lat_ms")
    if apply_lat and apply_lat > 50:
        recs.append(("⚠️ OSD 延迟", f"OSD 平均 Apply 延迟 {apply_lat:.1f} ms（建议 < 50 ms），请检查磁盘 IO 负载，考虑添加 SSD 缓存层。"))
    if commit_lat and commit_lat > 200:
        recs.append(("⚠️ OSD 延迟", f"OSD 平均 Commit 延迟 {commit_lat:.1f} ms 偏高，请检查后端网络和存储性能。"))
 
    # MON
    if mon["total"] < 3:
        recs.append(("⚠️ MON 状态", f"MON 节点仅 {mon['total']} 个，少于推荐的 3 个，存在单点故障风险，建议扩充至奇数个 MON。"))
    if mon["out_quorum"] > 0:
        recs.append(("❌ MON 状态", f"有 {mon['out_quorum']} 个 MON 未加入 Quorum，请检查 MON 服务状态（ceph mon stat）。"))
    if mon["total"] >= 3 and mon["out_quorum"] == 0:
        recs.append(("✅ MON 状态", f"所有 {mon['total']} 个 MON 均在 Quorum 中，运行正常。"))
 
    # PG
    degraded  = int(pg.get("degraded") or 0)
    stale     = int(pg.get("stale") or 0)
    undersized = int(pg.get("undersized") or 0)
    clean     = int(pg.get("clean") or 0)
    total_pg  = int(pg.get("total") or 0)
    if stale > 0:
        recs.append(("❌ PG 状态", f"有 {stale} 个 PG 处于 Stale 状态，可能导致 IO 阻塞，请立即排查！执行 `ceph pg dump_stuck stale`。"))
    if degraded > 0:
        recs.append(("⚠️ PG 状态", f"有 {degraded} 个 PG 处于 Degraded 状态，数据副本不完整，请关注 OSD 恢复进度（`ceph -s`）。"))
    if undersized > 0:
        recs.append(("⚠️ PG 状态", f"有 {undersized} 个 PG 处于 Undersized 状态，建议检查副本配置和 OSD 数量。"))
    if total_pg > 0 and clean == total_pg:
        recs.append(("✅ PG 状态", f"全部 {total_pg} 个 PG 处于 active+clean 状态，数据完整。"))
 
    # 性能
    rd_lat = perf.get("rd_lat_ms")
    wr_lat = perf.get("wr_lat_ms")
    if rd_lat and rd_lat > 20:
        recs.append(("⚠️ 读延迟", f"集群读延迟 {rd_lat:.1f} ms 偏高（建议 < 10 ms），请检查网络、OSD 负载和缓存命中率。"))
    if wr_lat and wr_lat > 30:
        recs.append(("⚠️ 写延迟", f"集群写延迟 {wr_lat:.1f} ms 偏高（建议 < 15 ms），请检查 Journal/WAL 配置和磁盘性能。"))
 
    # 告警
    if alts:
        by_sev = {}
        for a in alts:
            s = a.get("severity", "unknown")
            by_sev[s] = by_sev.get(s, 0) + 1
        sev_str = "、".join(f"{k} {v} 条" for k, v in by_sev.items())
        recs.append(("❌ 活跃告警", f"当前有 {len(alts)} 条活跃告警（{sev_str}），请按优先级逐一处理，高级别告警须在 1 小时内响应。"))
    else:
        recs.append(("✅ 告警信息", "当前 Alertmanager 无活跃告警。"))
 
    return recs
 
# ═══════════════════════════════════════════════════════════════
#  报告页面：封面
# ═══════════════════════════════════════════════════════════════
 
def build_cover(doc: Document, report_date: str, inspector: str):
    for _ in range(4):
        doc.add_paragraph()
 
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Ceph 存储集群自动化巡检报告")
    r.bold = True; r.font.size = Pt(26); r.font.name = "微软雅黑"
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    try:
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass
 
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Pacific 16.2.x ｜ cephadm + Docker ｜ 共 4 套集群")
    rs.font.size = Pt(13); rs.font.name = "微软雅黑"
    rs.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
 
    doc.add_paragraph()
    doc.add_paragraph()
 
    info_lines = [
        f"巡检日期：{report_date}",
        f"巡检人员：{inspector}",
        "报告版本：v1.0",
        "保密等级：内部文件",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p.add_run(line)
        r2.font.size = Pt(12); r2.font.name = "微软雅黑"
        r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
 
    doc.add_page_break()
 
 
# ═══════════════════════════════════════════════════════════════
#  报告页面：总览汇总表
# ═══════════════════════════════════════════════════════════════
 
def build_summary(doc: Document, all_data: List[Dict]):
    add_h(doc, "一、巡检总览", level=1)
    add_p(doc, "本次巡检覆盖 4 套 Ceph Pacific 集群，汇总如下：", size=10)
 
    headers = ["集群名称", "运行环境", "健康状态", "总容量",
               "使用率", "OSD总数", "DOWN", "MON数", "PG异常", "活跃告警"]
    table = make_table(doc, headers)
 
    for i, d in enumerate(all_data):
        if not d.get("reachable"):
            add_row(table, [d["cluster_info"]["name"], "-", "不可达", "-", "-", "-", "-", "-", "-", "-"],
                    center_cols=list(range(2, 10)), alt=(i % 2 == 0), row_color="FFD7D7")
            continue
        h   = d["health"]
        cap = d["capacity"]
        osd = d["osd"]
        mon = d["mon"]
        pg  = d["pg"]
        alt = d["alerts"]
        degraded = int(pg.get("degraded") or 0) + int(pg.get("stale") or 0)
        hcolor = None
        if h["code"] == "HEALTH_ERR":
            hcolor = "FFD7D7"
        elif h["code"] == "HEALTH_WARN":
            hcolor = THEME_YELL
        add_row(table,
            [d["cluster_info"]["name"], d["cluster_info"].get("env", "-"),
             h["code"], bh(cap.get("total")), fp(cap.get("usage_pct")),
             str(osd["total"]), str(osd["down"]), str(mon["total"]),
             str(degraded), str(len(alt))],
            center_cols=list(range(2, 10)), alt=(i % 2 == 0),
            row_color=hcolor,
        )
 
    doc.add_paragraph()
    doc.add_page_break()
 
 
# ═══════════════════════════════════════════════════════════════
#  报告页面：单集群巡检章节
# ═══════════════════════════════════════════════════════════════
 
CHAPTER_NUMS = ["二", "三", "四", "五", "六", "七", "八", "九"]
 
def build_cluster_section(doc: Document, data: Dict, chapter_idx: int):
    cfg  = data["cluster_info"]
    h    = data["health"]
    cap  = data["capacity"]
    osd  = data["osd"]
    mon  = data["mon"]
    pg   = data["pg"]
    pool = data["pool"]
    perf = data["performance"]
    alts = data["alerts"]
    cn   = CHAPTER_NUMS[chapter_idx] if chapter_idx < len(CHAPTER_NUMS) else str(chapter_idx + 2)
 
    # ─ 章标题
    add_h(doc, f"{cn}、{cfg['name']} 巡检报告", level=1)
 
    # ─ 1. 基本信息
    add_h(doc, "1. 集群基本信息", level=2)
    t = make_table(doc, ["配置项", "配置值"], THEME_DARK)
    rows = [
        ("集群名称",    cfg["name"]),
        ("运行环境",    cfg.get("env", "生产环境")),
        ("Ceph 版本",   cfg.get("version", "16.2.x")),
        ("部署方式",    "cephadm + Docker"),
        ("Prometheus",  cfg["prometheus_url"]),
        ("集群描述",    cfg.get("description", "-")),
        ("巡检时间",    data["collected_at"]),
    ]
    for i, (k, v) in enumerate(rows):
        add_row(t, [k, v], alt=(i % 2 == 0))
    doc.add_paragraph()
 
    # ─ 2. 健康状态
    add_h(doc, "2. 集群健康状态", level=2)
    t = make_table(doc, ["指标", "英文状态", "中文说明"], THEME_DARK)
    hc = HEALTH_COLOR.get(h["code"], "808080")
    row = t.add_row()
    for i, val in enumerate(["集群健康状态", h["code"], h["cn"]]):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        _set_font(run, size=10, bold=(i > 0),
                  color=tuple(int(hc[j:j+2], 16) for j in (0, 2, 4)) if i > 0 else None)
        _cell_border(cell)
    doc.add_paragraph()
 
    # ─ 3. 存储容量
    add_h(doc, "3. 存储容量分析", level=2)
    pct = cap.get("usage_pct")
    if pct and pct >= 85:
        cap_note, note_color = f"❌ 容量已达 {pct:.1f}%，超过 85% 警戒线，请立即扩容！", (0xC0, 0x00, 0x00)
    elif pct and pct >= 70:
        cap_note, note_color = f"⚠️ 容量 {pct:.1f}%，超 70% 预警线，关注增长趋势", (0xC0, 0x80, 0x00)
    else:
        cap_note, note_color = f"✅ 容量 {fp(pct)}，处于健康范围", (0x00, 0x80, 0x00)
 
    t = make_table(doc, ["指标", "数值"], THEME_DARK)
    cap_rows = [
        ("集群总容量",  bh(cap.get("total"))),
        ("已使用容量",  bh(cap.get("used"))),
        ("可用容量",    bh(cap.get("avail"))),
        ("使用率",      fp(pct)),
        ("容量评估",    cap_note),
    ]
    for i, (k, v) in enumerate(cap_rows):
        add_row(t, [k, v], center_cols=[1], alt=(i % 2 == 0))
    doc.add_paragraph()
 
    # ─ 4. OSD 状态
    add_h(doc, "4. OSD 状态分析", level=2)
    add_h(doc, "4.1 OSD 汇总", level=3)
    t = make_table(doc, ["指标", "数值"], THEME_DARK)
    osd_sum = [
        ("OSD 总数",        str(osd["total"])),
        ("UP 状态",         str(osd["up"])),
        ("DOWN 状态",       str(osd["down"])),
        ("IN 状态",         str(osd["in"])),
        ("OUT 状态",        str(osd["out"])),
        ("平均 Apply 延迟", fv(osd.get("avg_apply_lat_ms"), " ms")),
        ("平均 Commit 延迟",fv(osd.get("avg_commit_lat_ms"), " ms")),
    ]
    for i, (k, v) in enumerate(osd_sum):
        rcolor = None
        if k == "DOWN 状态" and osd["down"] > 0:
            rcolor = "FFD7D7"
        add_row(t, [k, v], center_cols=[1], alt=(i % 2 == 0), row_color=rcolor)
    doc.add_paragraph()
 
    if osd["details"]:
        add_h(doc, "4.2 OSD 详情", level=3)
        t = make_table(doc, ["OSD ID", "UP", "IN", "总容量", "已用", "使用率"])
        for i, o in enumerate(osd["details"]):
            is_down = o["up"] != 1
            add_row(t,
                [str(o["id"]),
                 "UP" if o["up"] == 1 else "DOWN",
                 "IN" if o["in"] == 1 else "OUT",
                 bh(o["total_bytes"]), bh(o["used_bytes"]), fp(o["usage_pct"])],
                center_cols=[0, 1, 2, 5],
                alt=(i % 2 == 0),
                row_color="FFD7D7" if is_down else None,
            )
        doc.add_paragraph()
 
    # ─ 5. MON 状态
    add_h(doc, "5. MON 状态分析", level=2)
    t = make_table(doc, ["指标", "数值"], THEME_DARK)
    mon_sum = [
        ("MON 总数",        str(mon["total"])),
        ("参与 Quorum",     str(mon["in_quorum"])),
        ("未参与 Quorum",   str(mon["out_quorum"])),
    ]
    for i, (k, v) in enumerate(mon_sum):
        rcolor = "FFD7D7" if k == "未参与 Quorum" and mon["out_quorum"] > 0 else None
        add_row(t, [k, v], center_cols=[1], alt=(i % 2 == 0), row_color=rcolor)
    doc.add_paragraph()
 
    if mon["details"]:
        add_h(doc, "5.1 MON 节点详情", level=3)
        t = make_table(doc, ["MON 节点名称", "Quorum 状态"])
        for i, m in enumerate(mon["details"]):
            in_q = m["in_quorum"] == 1
            add_row(t,
                [str(m["name"]), "✔ 在线" if in_q else "✘ 离线"],
                center_cols=[1], alt=(i % 2 == 0),
                row_color=None if in_q else "FFD7D7",
            )
        doc.add_paragraph()
 
    # ─ 6. PG 状态
    add_h(doc, "6. PG 状态分析", level=2)
    t = make_table(doc, ["PG 状态", "数量", "是否正常"], THEME_DARK)
    pg_rows = [
        ("total",        "PG 总数",    lambda v: "–"),
        ("active",       "Active",     lambda v: "✔ 正常" if v and v > 0 else "–"),
        ("clean",        "Clean",      lambda v: "✔ 正常" if v and v > 0 else "–"),
        ("degraded",     "Degraded",   lambda v: "❌ 异常" if v and v > 0 else "✔ 正常"),
        ("undersized",   "Undersized", lambda v: "⚠️ 注意" if v and v > 0 else "✔ 正常"),
        ("recovering",   "Recovering", lambda v: "⚠️ 恢复中" if v and v > 0 else "✔ 正常"),
        ("recovery_wait","Recovery Wait",lambda v:"⚠️ 等待" if v and v > 0 else "✔ 正常"),
        ("stale",        "Stale",      lambda v: "❌ 异常" if v and v > 0 else "✔ 正常"),
        ("peering",      "Peering",    lambda v: "⚠️ 注意" if v and v > 0 else "✔ 正常"),
        ("remapped",     "Remapped",   lambda v: "⚠️ 注意" if v and v > 0 else "✔ 正常"),
    ]
    for i, (key, label, status_fn) in enumerate(pg_rows):
        v = pg.get(key)
        cnt = fv(v, d=0) if v is not None else "N/A"
        status = status_fn(v)
        is_bad = "❌" in status or ("⚠️" in status and v and v > 0)
        add_row(t, [label, cnt, status],
                center_cols=[1, 2], alt=(i % 2 == 0),
                row_color="FFD7D7" if "❌" in status and v and v > 0 else
                          THEME_YELL if is_bad and "⚠️" in status else None)
    doc.add_paragraph()
 
    # ─ 7. Pool 状态
    add_h(doc, "7. Pool 状态分析", level=2)
    pools = pool.get("pools", [])
    if pools:
        t = make_table(doc,
            ["Pool 名称", "已用", "可用", "使用率", "对象数",
             "读 IOPS", "写 IOPS", "读带宽", "写带宽"])
        for i, p in enumerate(pools):
            pname = p["name"] or f"pool-{p['id']}"
            add_row(t,
                [pname,
                 bh(p["used"]), bh(p["avail"]), fp(p["usage_pct"]),
                 fv(p["objects"], d=0),
                 fv(p["rd_iops"], " ops/s", 1), fv(p["wr_iops"], " ops/s", 1),
                 bh(p["rd_bytes_s"]) + "/s", bh(p["wr_bytes_s"]) + "/s"],
                center_cols=list(range(1, 9)), alt=(i % 2 == 0),
                row_color="FFD7D7" if p["usage_pct"] >= 85 else
                          THEME_YELL if p["usage_pct"] >= 70 else None,
            )
    else:
        add_p(doc, "未能获取 Pool 数据，请检查 ceph_pool_bytes_used 指标。", size=9)
    doc.add_paragraph()
 
    # ─ 8. 性能指标
    add_h(doc, "8. 集群性能指标（当前）", level=2)
    t = make_table(doc, ["性能指标", "当前值", "参考阈值", "状态"], THEME_DARK)
    perf_rows = [
        ("读 IOPS",    fv(perf.get("rd_iops"), " ops/s", 1), "-",      "–"),
        ("写 IOPS",    fv(perf.get("wr_iops"), " ops/s", 1), "-",      "–"),
        ("读吞吐量",   bh(perf.get("rd_bytes_s")) + "/s" if perf.get("rd_bytes_s") else "N/A", "-", "–"),
        ("写吞吐量",   bh(perf.get("wr_bytes_s")) + "/s" if perf.get("wr_bytes_s") else "N/A", "-", "–"),
        ("读延迟",     fv(perf.get("rd_lat_ms"), " ms"), "< 10 ms",
         "⚠️ 偏高" if perf.get("rd_lat_ms") and perf["rd_lat_ms"] > 10 else "✔ 正常"),
        ("写延迟",     fv(perf.get("wr_lat_ms"), " ms"), "< 15 ms",
         "⚠️ 偏高" if perf.get("wr_lat_ms") and perf["wr_lat_ms"] > 15 else "✔ 正常"),
    ]
    for i, row_vals in enumerate(perf_rows):
        is_warn = "⚠️" in row_vals[3]
        add_row(t, list(row_vals), center_cols=[1, 2, 3], alt=(i % 2 == 0),
                row_color=THEME_YELL if is_warn else None)
    doc.add_paragraph()
 
    # ─ 9. 告警信息
    add_h(doc, "9. 当前活跃告警", level=2)
    if alts:
        t = make_table(doc, ["告警名称", "级别", "摘要", "触发时间"], THEME_RED)
        for i, a in enumerate(alts):
            is_crit = a.get("severity", "").lower() == "critical"
            add_row(t,
                [a["name"], a["severity"], a["summary"], a["fired_at"]],
                alt=(i % 2 == 0),
                row_color="FFD7D7" if is_crit else THEME_YELL,
            )
    else:
        add_p(doc, "✅  当前无活跃告警。", size=10, color=(0x00, 0x80, 0x00))
    doc.add_paragraph()
 
    # ─ 10. 巡检结论与建议
    add_h(doc, "10. 巡检结论与建议", level=2)
    recs = make_recommendations(data)
    t = make_table(doc, ["类别", "建议内容"], THEME_DARK)
    for i, (label, content) in enumerate(recs):
        add_row(t, [label, content], alt=(i % 2 == 0))
    doc.add_paragraph()
    doc.add_page_break()
 
 
# ═══════════════════════════════════════════════════════════════
#  报告主入口
# ═══════════════════════════════════════════════════════════════
 
def build_report(all_data: List[Dict], output_path: str, inspector: str):
    doc = Document()
 
    # A4 页面设置
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.54)
    section.top_margin  = section.bottom_margin = Cm(2.54)
 
    # 默认正文字体
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10)
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass
 
    report_date = datetime.datetime.now().strftime("%Y 年 %m 月 %d 日")
    build_cover(doc, report_date, inspector)
    build_summary(doc, all_data)
 
    for idx, data in enumerate(all_data):
        build_cluster_section(doc, data, idx)
 
    doc.save(output_path)
 
 
# ═══════════════════════════════════════════════════════════════
#  并行采集 + CLI
# ═══════════════════════════════════════════════════════════════
 
def collect_cluster(cfg: Dict) -> Dict:
    try:
        return CephCollector(cfg).collect()
    except Exception as exc:
        print(f"  [ERROR] {cfg['name']} 采集异常: {exc}")
        return {
            "cluster_info": cfg,
            "collected_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "reachable": False,
            "health":    {"value": None, "code": "UNKNOWN", "cn": "采集失败"},
            "capacity":  {"total": None, "used": None, "avail": None, "usage_pct": None},
            "osd":       {"total": 0, "up": 0, "in": 0, "down": 0, "out": 0,
                          "details": [], "avg_apply_lat_ms": None, "avg_commit_lat_ms": None},
            "mon":       {"total": 0, "in_quorum": 0, "out_quorum": 0, "details": []},
            "pg":        {}, "pool": {"pools": []}, "performance": {}, "alerts": [],
        }
 
 
def main():
    parser = argparse.ArgumentParser(description="Ceph 集群自动化巡检工具")
    parser.add_argument("--output", default=REPORT_OUTPUT_DIR,
                        help=f"报告输出目录 (默认: {REPORT_OUTPUT_DIR})")
    parser.add_argument("--cluster", default=None,
                        help="只巡检指定名称的集群（模糊匹配）")
    parser.add_argument("--inspector", default="运维团队",
                        help="巡检人员姓名/团队名")
    parser.add_argument("--parallel", action="store_true", default=True,
                        help="并行采集多集群（默认启用）")
    args = parser.parse_args()
 
    print("=" * 62)
    print("   Ceph 集群自动化巡检工具  v1.0")
    print("   Pacific 16.2.x | cephadm + Docker")
    print("=" * 62)
 
    targets = CLUSTERS
    if args.cluster:
        targets = [c for c in CLUSTERS if args.cluster in c["name"]]
        if not targets:
            print(f"[ERROR] 未找到匹配集群: {args.cluster}")
            sys.exit(1)
 
    print(f"\n[INFO] 共 {len(targets)} 套集群待巡检，并行模式: {args.parallel}")
 
    all_data: List[Dict] = [None] * len(targets)
 
    if args.parallel and len(targets) > 1:
        with ThreadPoolExecutor(max_workers=len(targets)) as pool:
            futures = {pool.submit(collect_cluster, cfg): i
                       for i, cfg in enumerate(targets)}
            for fut in as_completed(futures):
                idx = futures[fut]
                all_data[idx] = fut.result()
    else:
        for i, cfg in enumerate(targets):
            all_data[i] = collect_cluster(cfg)
 
    os.makedirs(args.output, exist_ok=True)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out_file = os.path.join(args.output, f"ceph_inspection_{ts}.docx")
 
    print(f"\n[INFO] 正在生成 Word 巡检报告 ...")
    build_report(all_data, out_file, args.inspector)
 
    print(f"\n{'=' * 62}")
    print(f"  ✅  巡检完成！")
    print(f"  📄  报告路径: {out_file}")
    print(f"{'=' * 62}\n")
 
 
if __name__ == "__main__":
    main()
